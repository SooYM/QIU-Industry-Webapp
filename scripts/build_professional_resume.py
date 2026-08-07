from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "resume" / "Soo_Yau_Ming_Professional_Resume.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

FONT = "Arial"
NAVY = RGBColor(24, 59, 78)
TEAL = RGBColor(45, 111, 115)
INK = RGBColor(28, 32, 35)
MUTED = RGBColor(92, 101, 107)


def font(run, size, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def border_bottom(paragraph, color="2D6F73", size="12", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def shade(paragraph, fill="EDF5F4"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def keep_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def section_heading(doc, text):
    p = doc.add_paragraph(style="Professional Section")
    p.add_run(text.upper())
    border_bottom(p, color="B8CCCF", size="5", space="2")
    keep_next(p)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="Professional Bullet")
    r = p.add_run(text)
    font(r, 9.55)
    return p


def label_line(doc, label, value):
    p = doc.add_paragraph(style="Professional Body")
    p.paragraph_format.space_after = Pt(1.8)
    a = p.add_run(label + "  ")
    font(a, 9.2, bold=True, color=NAVY)
    b = p.add_run(value)
    font(b, 9.2)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.58)
    sec.bottom_margin = Inches(0.58)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04

    body = doc.styles.add_style("Professional Body", WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    body.font.name = FONT
    body._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    body._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    body.font.size = Pt(9.6)
    body.paragraph_format.space_after = Pt(4)
    body.paragraph_format.line_spacing = 1.05

    heading = doc.styles.add_style("Professional Section", WD_STYLE_TYPE.PARAGRAPH)
    heading.base_style = normal
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading.font.size = Pt(10.2)
    heading.font.bold = True
    heading.font.color.rgb = NAVY
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.line_spacing = 1.0

    bullet_style = doc.styles.add_style("Professional Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.base_style = doc.styles["List Bullet"]
    bullet_style.font.name = FONT
    bullet_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet_style.font.size = Pt(9.55)
    bullet_style.paragraph_format.left_indent = Inches(0.24)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.17)
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.space_after = Pt(3.2)
    bullet_style.paragraph_format.line_spacing = 1.03

    name = doc.add_paragraph()
    name.paragraph_format.space_after = Pt(0)
    r = name.add_run("SOO YAU MING")
    font(r, 25, bold=True, color=NAVY)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("SOFTWARE ENGINEER  |  APPLIED AI & COMPUTER VISION")
    font(r, 10.8, bold=True, color=TEAL)

    contact = doc.add_paragraph()
    contact.paragraph_format.space_after = Pt(8)
    r = contact.add_run("Ipoh, Perak, Malaysia  |  [PHONE]  |  [PROFESSIONAL EMAIL]  |  [LINKEDIN]  |  [GITHUB]")
    font(r, 8.8, color=MUTED)
    border_bottom(contact, color="2D6F73", size="13", space="5")

    section_heading(doc, "Profile")
    p = doc.add_paragraph(style="Professional Body")
    r = p.add_run(
        "Computer Science undergraduate and Research Assistant Intern building full-stack products and applied AI systems. Experience spans role-based web platforms, healthcare OCR, cloud data workflows, and computer vision model development, with regular delivery to executive, academic, and operational stakeholders."
    )
    font(r, 9.7)

    impact = doc.add_paragraph()
    impact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    impact.paragraph_format.left_indent = Inches(0.06)
    impact.paragraph_format.right_indent = Inches(0.06)
    impact.paragraph_format.space_before = Pt(3)
    impact.paragraph_format.space_after = Pt(5)
    impact.paragraph_format.line_spacing = 1.0
    shade(impact)
    r = impact.add_run("40+ PORTAL IMPROVEMENTS    |    135,569 ML IMAGES    |    78.77% VALIDATION ACCURACY")
    font(r, 9.2, bold=True, color=NAVY)

    section_heading(doc, "Experience")
    role = doc.add_paragraph()
    role.paragraph_format.space_after = Pt(0)
    r = role.add_run("Research Assistant Intern")
    font(r, 11, bold=True, color=NAVY)
    r = role.add_run("  |  Quest International University - Delta Lab  |  Apr 2026 - Present")
    font(r, 9.5, bold=True)
    keep_next(role)

    place = doc.add_paragraph(style="Professional Body")
    place.paragraph_format.space_after = Pt(3)
    r = place.add_run("Ipoh, Malaysia")
    font(r, 9, italic=True, color=MUTED)
    keep_next(place)

    bullet(
        doc,
        "Shipped 40+ technical and UI improvements in one development cycle for QIU Industry Day 2026 by building a Next.js, React, TypeScript, and Firebase career platform with 4-tier access control, role-scoped dashboards, employer approvals, analytics, CSV exports, and automated CV generation.",
    )
    bullet(
        doc,
        "Improved career discovery with degree-aware vacancy recommendations, a 5-mode sorting engine, application tracking, and a grounded RAG assistant that answers from company and job-profile data.",
    )
    bullet(
        doc,
        "Scaled a 20-species aerial classification pipeline to 135,569 images and raised validation accuracy from 76.39% to 78.20% through 25-epoch EfficientNet-B3 training; later reached 78.77% with EfficientNetV2-S and 25 augmented subsets.",
    )
    bullet(
        doc,
        "Validated a Flutter healthcare OCR workflow against 48 synthetic blood reports spanning 12 months and 4 patient profiles by creating a Python/WeasyPrint test generator and integrating Google ML Kit preprocessing with OpenAI API extraction.",
    )
    bullet(
        doc,
        "Launched a domain-restricted judging and voting portal after reaching 100% evaluation-panel submission compliance by auditing scores, enforcing OAuth/SSO access, deduplicating ballots, and applying deadline-based database locks.",
    )

    section_heading(doc, "Technical Toolkit")
    label_line(doc, "Languages", "TypeScript, JavaScript, Python, Dart, SQL, HTML, CSS")
    label_line(doc, "Product Engineering", "Next.js, React, Flutter, Tailwind CSS, Firebase Authentication, Cloud Firestore, REST APIs")
    label_line(doc, "AI & Data", "Computer Vision, OCR, EfficientNet, Mask R-CNN, RAG, Google ML Kit, OpenAI API, Gemini API, BigQuery, Supabase")

    section_heading(doc, "Education & Leadership")
    education = doc.add_paragraph(style="Professional Body")
    education.paragraph_format.space_after = Pt(2)
    r = education.add_run("Bachelor of Computer Science")
    font(r, 10, bold=True, color=NAVY)
    r = education.add_run("  |  Quest International University  |  Expected [MONTH YEAR]")
    font(r, 9.4, bold=True)

    p = doc.add_paragraph(style="Professional Body")
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "Co-facilitated an introductory machine learning practical and presented full-stack and AI prototypes to the COO, Career and Professional Development Centre, lecturers, and research stakeholders."
    )
    font(r, 9.4)

    doc.core_properties.title = "Soo Yau Ming - Professional Resume"
    doc.core_properties.subject = "Software Engineering and Applied AI Resume"
    doc.core_properties.author = "Soo Yau Ming"
    doc.core_properties.keywords = "Software Engineer, AI, Machine Learning, Next.js, React, TypeScript, Flutter, Python, Firebase, Computer Vision, OCR"
    doc.core_properties.comments = "Professional single-column resume built from verified industrial training evidence."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
